from docx import Document
from datetime import datetime
from math import ceil


class CreateDeclaration:
    def __init__(self, parking_place, date, name_yacht, registration_number, home_port, yacht_length, yacht_width,
                 yacht_type, owner_details, commissioning_body, parking_peroid, chip_card):
        self.document = Document('deklaracja.docx')
        self.parking_peroid = parking_peroid
        self.days = self.count_days()
        self.parking_place = parking_place
        self.date = date
        self.name_yacht = name_yacht
        self.registration_number = registration_number
        self.home_port = home_port
        self.yacht_length = yacht_length
        self.yacht_width = yacht_width
        self.yacht_type = yacht_type
        self.owner_details = owner_details
        self.commissioning_body = commissioning_body

        self.parking_fee = round(self.parking_fee(), 2)
        self.quarter_fee = round(self.parking_fee/4.0, 2)
        self.chip_card = chip_card


    def count_days(self):
        # Obliczenie ilości dni w okresie postoju
        date_format = "%d.%m.%Y"
        a = datetime.strptime(self.parking_peroid['from'], date_format)
        b = datetime.strptime(self.parking_peroid['to'], date_format)
        delta = b - a
        return delta.days

    def summer_fee(self):
        # obliczenie opłaty za sezon letni
        # Jeżeli okres dłuższy niż 184 dni to przyszługuje zniżka 30%
        if self.days >= 183:
            return ceil(self.yacht_length) * 4.30 * 184 * 0.7
        else:
            return ceil(self.yacht_length) * 4.30 * self.days

    def winter_fee(self):
        return ceil(self.yacht_length) * ceil(self.yacht_width) * 0.25

    def parking_fee(self):
        if self.days in [364, 365, 366]:
            return self.summer_fee() + self.winter_fee()
        else:
            return self.summer_fee()

    def create_document(self):
        for paragraph in self.document.paragraphs:
            if 'Miejsce postojowe' in paragraph.text:
                paragraph.text = 'Miejsce postojowe {}                                                                  ' \
                                 '                                        ' \
                                 'Data {}'.format(self.parking_place, self.date)

        self.document.add_paragraph('\n')
        self.document.add_paragraph('Nazwa jachtu: ', style='List Number').add_run(self.name_yacht).bold = True
        self.document.add_paragraph('Nr rejestracyjny: ', style='List Number').add_run(
            self.registration_number).bold = True
        self.document.add_paragraph('Port macierzysty: ', style='List Number').add_run(self.home_port).bold = True
        self.document.add_paragraph('Dane jachtu: ', style='List Number')
        p = self.document.add_paragraph('            Długość: ')
        p.add_run(str(self.yacht_length)).bold = True
        p.add_run('             Szerokość: ')
        p.add_run(str(self.yacht_width)).bold = True
        self.document.add_paragraph('Typ jachtu: ', style='List Number').add_run(self.yacht_type).bold = True
        self.document.add_paragraph('Dane właściciela* lub użytkownika jachtu*: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Dane armatora: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Podmiot zlecający, podpisujący umowę na postój jachtu: ', style='List Number')
        self.document.add_paragraph('            - nazwisko imię */ pełna nazwa klubu lub firmy: ').add_run(
            self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres: ').add_run(self.commissioning_body['address']).bold = True
        self.document.add_paragraph('            - tel: ').add_run(self.commissioning_body['tel']).bold = True
        self.document.add_paragraph('            - E-mial: ').add_run(self.commissioning_body['e-mail']).bold = True
        self.document.add_paragraph('            - NIP klubu/stowarzyszenia: ').add_run(
            self.commissioning_body['nip']).bold = True
        self.document.add_paragraph('Deklarowany okres i czas postoju: ', style='List Number')
        self.document.add_paragraph('            - od dnia: ').add_run(self.parking_peroid['from']).bold = True
        self.document.add_paragraph('            - do dnia: ').add_run(self.parking_peroid['to']).bold = True
        self.document.add_paragraph('')
        self.document.add_paragraph('Opłata za postój wynosi: ', style='List Number').add_run(
            str(self.parking_fee)).bold = True
        self.document.add_paragraph(
            'Czynsz płatny, zgodnie z wystawioną fakturą z góry, jednorazowo lub w czterech poniższych '
            'ratach za każdy kwartał do:')
        self.document.add_paragraph('      I. 15.05.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('     II. 15.08.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('    III. 15.11.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('    IV. 15.21.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('na rachunek Wynajmującego o nr  88 1030 1117 0000 0000 8899 5007.')
        self.document.add_paragraph('Adres do korespondencji: ', style='List Number').add_run(
            self.owner_details['address']).bold = True
        self.document.add_paragraph('Karta chipowa: ', style='List Number').add_run(self.chip_card).bold = True
        self.document.add_paragraph('Miejsce postoju nr: ', style='List Number').add_run(self.parking_place).bold = True
        self.document.add_paragraph(
            '\nUprzejmie informujemy, że w trakcie postoju jachtu na przystani NCŻ AWFiS może zaistnieć '
            'konieczność zmiany miejsca lokalizacji postoju jachtu na inne zgodnie z ze wskazaniem '
            'upoważnionego pracownika przystani.\n\n'
            'Oświadczam, iż zapoznałem się z Regulaminem przystani Narodowego Centrum Żeglarstwa AWFiS '
            'Gdańsku, w tym zawarcia umowy na postój. W pełni go akceptuję co poświadczam własnoręcznym '
            'podpisem pod Deklaracją postoju. Oświadczam, iż w przypadku okresu postoju krótszego niż '
            'zadeklarowany (wynikający m.in. ze sprzedaży jachtu) mam świadomość, że stawka będzie '
            'rekalkulowana do krótszego okresu postoju, zgodnie z obowiązującym cennikiem.\n\n').paragraph_format.alignment = 3
        self.document.add_paragraph(".........................................................				 .........."
                                    "...............................................").add_run(
            '   podpis pracownika NCŻ                                                                                 '
            '             podpis (imię, nazwisko)').italic = True

        self.document.add_page_break()

        self.document.save('Deklaracja {}.docx'.format(self.name_yacht))

"""
date = '01.04.2020'
parking_place = 'B-02'
name_yacht = 'Fordzik'
registration_number = 'PL-102'
home_port = 'Gdańsk'
yacht_length = 12.0
yacht_width = 3.0
yacht_type = 'motorowy'
owner_details = {'name': 'Adrian Bieliński', 'address': '80-298 Gdańsk, ul. Politechniczna 9/9'}
commissioning_body = {'name': 'Adrian Bieliński', 'address': '80-298 Gdańsk, ul. Politechniczna 9/9',
                      'tel': '510-494-063', 'e-mail': 'a.bielinski58@gmail.com', 'nip': 'BRAK'}
parking_peroid = {'from': '01.05.2020', 'to': '31.10.2020'}
correspondence_address = owner_details['address']
chip_card = '1 szt.'

test = CreateDeclaration(parking_place, date, name_yacht, registration_number, home_port, yacht_length, yacht_width,
                 yacht_type, owner_details, commissioning_body, parking_peroid, chip_card)
test.create_document()
"""